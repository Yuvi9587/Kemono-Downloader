import os
import queue
import re
import threading
import time
import sys
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from urllib.parse import urlparse

import cloudscraper
import requests
from curl_cffi import requests as cffi_requests
from PySide6.QtCore import QThread, Signal

from PIL import Image
try:
    import imagehash
    IMAGEHASH_AVAILABLE = True
except ImportError:
    imagehash = None
    IMAGEHASH_AVAILABLE = False
from ...core.database_manager import DatabaseManager

from ...core.bunkr_client import fetch_bunkr_data
from ...core.pixeldrain_client import fetch_pixeldrain_data
from ...core.saint2_client import fetch_saint2_data
from ...core.simpcity_client import fetch_single_simpcity_page
from ...services.drive_downloader import (
    download_mega_file as drive_download_mega_file
)
from ...utils.file_utils import clean_folder_name, clean_filename, format_custom_suffix, format_custom_date
from ...core.workers import robust_clean_name
from ...ui.dialogs.SinglePDF import create_individual_pdf, create_single_pdf_from_content
from ...utils.proxy_utils import get_proxies_from_settings

try:
    from docx import Document
except ImportError:
    Document = None


class SimpCityDownloadThread(QThread):
    progress_signal = Signal(str)
    file_progress_signal = Signal(str, object)
    finished_signal = Signal(int, int, bool, list)
    overall_progress_signal = Signal(int, int)

    def __init__(self, url, post_id, output_dir, cookies, parent=None):
        super().__init__(parent)
        self.start_url = url
        self.post_id = post_id
        self.output_dir = output_dir
        self.cookies = cookies
        self.is_cancelled = False
        self.parent_app = parent
        self.image_queue = queue.Queue()
        self.service_queue = queue.Queue()
        self.counter_lock = threading.Lock()
        self.total_dl_count = 0
        self.total_skip_count = 0
        self.total_jobs_found = 0
        self.total_jobs_processed = 0
        self.processed_job_urls = set()
        self.collected_external_links = []
        
        self.db = DatabaseManager()
        self.proxies = get_proxies_from_settings(parent.settings) if hasattr(parent, 'settings') else None

    def _check_pause_cancel(self):
        if self.is_cancelled or (self.parent_app and self.parent_app.cancellation_event.is_set()):
            self.is_cancelled = True
            return True
            
        if self.parent_app and self.parent_app.pause_event.is_set():
            self.progress_signal.emit("   Download paused...")
            while self.parent_app.pause_event.is_set():
                if self.is_cancelled or (self.parent_app and self.parent_app.cancellation_event.is_set()):
                    self.is_cancelled = True
                    return True
                time.sleep(0.5)
            self.progress_signal.emit("   Download resumed.")
            
        return self.is_cancelled

    def _smart_sleep(self, seconds):
        for _ in range(int(seconds * 10)):
            if self._check_pause_cancel(): return True
            time.sleep(0.1)
        return False

    def _record_to_db(self, filepath, filename):
        calculated_phash = None
        valid_exts = {'.jpg', '.jpeg', '.png', '.bmp', '.webp'}
        if os.path.splitext(filepath)[1].lower() in valid_exts:
            try:
                calculated_phash = str(imagehash.phash(Image.open(filepath), hash_size=16))
            except Exception:
                pass
        self.db.record_tagless_download(
            file_path=filepath,
            file_name=filename,
            file_hash=None,
            phash=calculated_phash
        )

    def cancel(self):
        self.is_cancelled = True

    class _ServiceLoggerAdapter:
        def __init__(self, signal_emitter, prefix=""):
            self.emit = signal_emitter
            self.prefix = prefix

        def __call__(self, msg, *args, **kwargs):
            self.info(msg, *args, **kwargs)
            
        def info(self, msg, *args, **kwargs): self.emit(f"{self.prefix}{str(msg) % args}")
        def error(self, msg, *args, **kwargs): self.emit(f"{self.prefix}❌ ERROR: {str(msg) % args}")
        def warning(self, msg, *args, **kwargs): self.emit(f"{self.prefix}⚠️ WARNING: {str(msg) % args}")

    def _log_interceptor(self, message):
        if "[SimpCity] Scraper found" in message or "[SimpCity] Scraping page" in message:
            pass
        else:
            self.progress_signal.emit(message)

    def _collect_external_links(self, jobs):
        if not getattr(self, 'external_links_config', None) or not jobs:
            return
        for job in jobs:
            if job.get('type') in ['bunkr', 'pixeldrain', 'mega', 'saint2', 'saint2_direct']:
                pm = job.get('post_metadata', {})
                self.collected_external_links.append({
                    "url": job.get('url', ''),
                    "platform": job.get('type', ''),
                    "thread_title": pm.get('thread_title', ''),
                    "post_id": str(pm.get('post_id', '')),
                    "username": pm.get('creator_name', ''),
                    "date": pm.get('published', '')
                })

    def _get_enriched_jobs(self, jobs_to_check):
        if not jobs_to_check:
            return []
            
        enriched_jobs = []
        bunkr_logger = self._ServiceLoggerAdapter(self.progress_signal.emit, prefix="      ")
        pixeldrain_logger = self._ServiceLoggerAdapter(self.progress_signal.emit, prefix="      ")
        saint2_logger = self._ServiceLoggerAdapter(self.progress_signal.emit, prefix="      ")
        
        for job in jobs_to_check:
            if self._check_pause_cancel(): break
            job_type = job.get('type')
            job_url = job.get('url')

            if job_type == 'image' and self.should_dl_images:
                enriched_jobs.append(job)
            elif job_type in ['saint2_direct', 'saint2'] and self.should_dl_saint2:
                enriched_jobs.append(job)
            elif job_type == 'mega' and self.should_dl_mega:
                enriched_jobs.append(job)
            elif (job_type == 'bunkr' and self.should_dl_bunkr) or \
                 (job_type == 'pixeldrain' and self.should_dl_pixeldrain):
                self.progress_signal.emit(f"   -> Checking {job_type} album for file count...")
                
                fetch_map = {
                    'bunkr': (fetch_bunkr_data, bunkr_logger),
                    'pixeldrain': (fetch_pixeldrain_data, pixeldrain_logger)
                }
                fetch_func, logger_adapter = fetch_map[job_type]
                album_name, files = fetch_func(job_url, logger_adapter, proxies=self.proxies)
                
                if files:
                    job['prefetched_files'] = files
                    job['prefetched_album_name'] = album_name
                    enriched_jobs.append(job)
        
        if enriched_jobs and not self.is_cancelled:
            summary_counts = Counter()
            current_page_file_count = 0
            for job in enriched_jobs:
                if job.get('prefetched_files'):
                    file_count = len(job['prefetched_files'])
                    summary_counts[job['type']] += file_count
                    current_page_file_count += file_count
                else:
                    summary_counts[job['type']] += 1
                    current_page_file_count += 1
            
            summary_parts = [f"{job_type} ({count})" for job_type, count in summary_counts.items()]
            self.progress_signal.emit(f"   [SimpCity] Content Found: {' | '.join(summary_parts)}")
            
            with self.counter_lock: self.total_jobs_found += current_page_file_count
            self.overall_progress_signal.emit(self.total_jobs_found, self.total_jobs_processed)

        return enriched_jobs

    def _apply_simpcity_custom_renaming(self, original_filename, post_metadata):
        if getattr(self.parent_app, 'manga_filename_style', None) != 'custom' or not post_metadata:
            return original_filename
            
        ext = os.path.splitext(original_filename)[1]
        custom_format = getattr(self.parent_app, 'custom_manga_filename_format', '{published} {creator_name} {id}')
        custom_suffix_format = getattr(self.parent_app, 'custom_manga_suffix_format', '001')
        
        custom_date_format = getattr(self.parent_app, 'manga_custom_date_format', 'YYYY-MM-DD')
        
        raw_published = post_metadata.get('published', 'Unknown Date')
        published = format_custom_date(raw_published, custom_date_format) if raw_published != 'Unknown Date' else 'Unknown Date'
        
        creator = robust_clean_name(post_metadata.get('thread_title', 'Unknown Thread'))
        post_id = str(post_metadata.get('post_id', 'unknown_id'))
        
        file_index = post_metadata.get('file_index', 1)
        suffix = format_custom_suffix(custom_suffix_format, file_index)
        
        raw_content = str(post_metadata.get('content', ''))
        content_no_newlines = raw_content.replace('\n', ' ').replace('\r', '')
        content = robust_clean_name(content_no_newlines).strip()[:50].strip()
        
        new_name = custom_format.replace('{published}', published).replace('{creator_name}', creator).replace('{id}', post_id).replace('{content}', content)
        
        has_suffix_placeholder = '{suffix}' in custom_format
        has_name_placeholder = '{name}' in custom_format
        
        if has_suffix_placeholder:
            new_name = new_name.replace('{suffix}', suffix)
            
        if has_name_placeholder:
            original_base, _ = os.path.splitext(original_filename)
            new_name = new_name.replace('{name}', original_base)
        
        new_name = robust_clean_name(new_name).strip()
        
        if not has_suffix_placeholder and not has_name_placeholder and new_name:
            new_name = f"{new_name}_{suffix}"
            
        return new_name + ext if new_name else original_filename

    def _get_target_path(self, base_album_path, post_metadata):
        if getattr(self, 'should_create_subfolder_per_post', False) and post_metadata:
            post_id = str(post_metadata.get('post_id', 'unknown'))
            published = post_metadata.get('published', 'Unknown Date')
            custom_date_format = getattr(self.parent_app, 'manga_custom_date_format', 'YYYY-MM-DD')
            formatted_date = format_custom_date(published, custom_date_format) if published != 'Unknown Date' else 'Unknown Date'
            subfolder_name = clean_folder_name(f"{post_id} - {formatted_date}")
            target_path = os.path.join(base_album_path, subfolder_name)
            os.makedirs(target_path, exist_ok=True)
            return target_path
        return base_album_path

    def _download_single_image(self, job, album_path, session):
        filename = job['filename']
        if 'post_metadata' in job:
            filename = self._apply_simpcity_custom_renaming(filename, job['post_metadata'])
            album_path = self._get_target_path(album_path, job['post_metadata'])
        filepath = os.path.join(album_path, filename)
        try:
            if os.path.exists(filepath):
                self.progress_signal.emit(f"   -> Skip (Image): '{filename}'")
                with self.counter_lock: self.total_skip_count += 1
                return
            self.progress_signal.emit(f"   -> Downloading (Image): '{filename}'...")
            response = session.get(job['url'], stream=True, timeout=180, headers={'Referer': self.start_url})
            response.raise_for_status()
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if self._check_pause_cancel():
                        f.close()
                        os.remove(filepath)
                        return
                    f.write(chunk)
            if not self.is_cancelled:
                self._record_to_db(filepath, filename)
                with self.counter_lock: self.total_dl_count += 1
        except Exception as e:
            self.progress_signal.emit(f"      -> ❌ Image download failed for '{filename}': {e}")
            with self.counter_lock: self.total_skip_count += 1
        finally:
            if not self.is_cancelled:
                with self.counter_lock: self.total_jobs_processed += 1
                self.overall_progress_signal.emit(self.total_jobs_found, self.total_jobs_processed)

    def _image_worker(self, album_path):
        session = cloudscraper.create_scraper()
        while True:
            if self._check_pause_cancel(): break
            try:
                job = self.image_queue.get(timeout=1)
                if job is None: break
                self._download_single_image(job, album_path, session)
                self.image_queue.task_done()
            except queue.Empty:
                continue

    def _service_worker(self, album_path):
        while True:
            if self._check_pause_cancel(): break
            try:
                job = self.service_queue.get(timeout=1)
                if job is None: break
                
                job_type = job['type']
                job_url = job['url']
                
                if job_type in ['pixeldrain', 'bunkr']:
                    if (job_type == 'pixeldrain' and self.should_dl_pixeldrain) or \
                       (job_type == 'bunkr' and self.should_dl_bunkr):
                        self.progress_signal.emit(f"\n--- Processing Service ({job_type.capitalize()}): {job_url} ---")
                        self._download_album(job.get('prefetched_files', []), job_url, album_path, job.get('post_metadata'))
                
                elif job_type in ['saint2', 'saint2_direct'] and self.should_dl_saint2:
                    self.progress_signal.emit(f"\n--- Processing Service (Saint2/Turbo): {job_url} ---")
                    saint2_logger = self._ServiceLoggerAdapter(self.progress_signal.emit, prefix="      ")
                    album_name, files = fetch_saint2_data(job_url, saint2_logger, proxies=self.proxies)
                    
                    if files:
                        if len(files) > 1:
                            with self.counter_lock:
                                self.total_jobs_found += (len(files) - 1)
                        self._download_album(files, job_url, album_path, job.get('post_metadata'))
                    else:
                        with self.counter_lock: 
                            self.total_jobs_processed += 1
                            self.overall_progress_signal.emit(self.total_jobs_found, self.total_jobs_processed)
                
                elif job_type == 'mega' and self.should_dl_mega:
                    self.progress_signal.emit(f"\n--- Processing Service (Mega): {job_url} ---")
                    target_path = self._get_target_path(album_path, job.get('post_metadata'))
                    drive_download_mega_file(job_url, target_path, self.progress_signal.emit, self.file_progress_signal.emit)
                    with self.counter_lock:
                        self.total_jobs_processed += 1
                        self.overall_progress_signal.emit(self.total_jobs_found, self.total_jobs_processed)
                
                self.service_queue.task_done()
            except queue.Empty:
                continue

    def _download_album(self, files_to_process, source_url, album_path, post_metadata=None):
        if not files_to_process: return
        
        session = cffi_requests.Session(impersonate="chrome120")
        
        if post_metadata:
            album_path = self._get_target_path(album_path, post_metadata)
            
        for file_data in files_to_process:
            if self._check_pause_cancel(): return 
            filename = file_data.get('filename') or file_data.get('name')
            if post_metadata:
                filename = self._apply_simpcity_custom_renaming(filename, post_metadata)
            filepath = os.path.join(album_path, filename)
            
            try:
                if os.path.exists(filepath):
                    self.progress_signal.emit(f"       -> Skip: '{filename}' already exists.")
                    with self.counter_lock: self.total_skip_count += 1
                else:
                    self.progress_signal.emit(f"       -> Downloading: '{filename}'...")
                    
                    headers = file_data.get('headers') or file_data.get('_http_headers') or {'Referer': source_url}
                    
                    if 'User-Agent' in headers:
                        headers = dict(headers)
                        headers.pop('User-Agent')
                        
                    req_cookies = file_data.get('cookies', {})
                    
                    response = session.get(file_data.get('url'), stream=True, timeout=180, headers=headers, cookies=req_cookies)
                    
                    content_type = response.headers.get('Content-Type', '')
                    if 'text/html' in content_type:
                        self.progress_signal.emit(f"       ❌ Blocked by CDN! Downloaded a 42KB HTML page instead of the video.")
                        with self.counter_lock: self.total_skip_count += 1
                        continue
                        
                    response.raise_for_status()
                    
                    total_size = int(response.headers.get('content-length', 0))
                    downloaded_size = 0
                    last_emit_time = 0
                    
                    with open(filepath, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if self._check_pause_cancel(): 
                                f.close()
                                os.remove(filepath)
                                return
                            f.write(chunk)
                            downloaded_size += len(chunk)
                            
                            current_time = time.time()
                            if current_time - last_emit_time > 0.5 or downloaded_size == total_size:
                                last_emit_time = current_time
                                if total_size > 0:
                                    percent = (downloaded_size / total_size) * 100
                                    self.file_progress_signal.emit(filename, f"[{percent:.1f}% - {downloaded_size/(1024*1024):.1f}MB/{total_size/(1024*1024):.1f}MB]")
                                else:
                                    self.file_progress_signal.emit(filename, f"[{downloaded_size/(1024*1024):.1f}MB]")
                    self.file_progress_signal.emit(None, None)
                            
                    if not self.is_cancelled:
                        self._record_to_db(filepath, filename)
                        with self.counter_lock: self.total_dl_count += 1
            except Exception as e:
                self.progress_signal.emit(f"       ❌ Download Error: {e}")
                with self.counter_lock: self.total_skip_count += 1
            finally:
                if not self.is_cancelled:
                    with self.counter_lock: self.total_jobs_processed += 1
                    self.overall_progress_signal.emit(self.total_jobs_found, self.total_jobs_processed)
    
    def run(self):
        self.progress_signal.emit("=" * 40)
        self.progress_signal.emit(f"🚀 Starting SimpCity Download for: {self.start_url}")

        self.should_dl_pixeldrain = self.parent_app.simpcity_dl_pixeldrain_cb.isChecked()
        self.should_dl_saint2 = self.parent_app.simpcity_dl_saint2_cb.isChecked()
        self.should_dl_mega = self.parent_app.simpcity_dl_mega_cb.isChecked()
        self.should_dl_images = self.parent_app.simpcity_dl_images_cb.isChecked()
        self.should_dl_bunkr = self.parent_app.simpcity_dl_bunkr_cb.isChecked()
        self.external_links_config = getattr(self.parent_app, 'simpcity_external_links_config', None)
        self.should_extract_text = getattr(self.parent_app, 'radio_more', None) and self.parent_app.radio_more.isChecked()
        self.should_create_subfolder_per_post = getattr(self.parent_app, 'use_subfolder_per_post_checkbox', None) and self.parent_app.use_subfolder_per_post_checkbox.isChecked()
        
        if self.should_extract_text:
            self.should_dl_pixeldrain = False
            self.should_dl_saint2 = False
            self.should_dl_mega = False
            self.should_dl_images = False
            self.should_dl_bunkr = False

        self.text_export_format = getattr(self.parent_app, 'text_export_format', 'txt')
        self.add_info_in_pdf_setting = False  # Enforced for SimpCity
        
        self.manga_mode_active = getattr(self.parent_app, 'manga_mode_checkbox', None) and self.parent_app.manga_mode_checkbox.isChecked()
        self.manga_custom_filename_format = getattr(self.parent_app, 'custom_manga_filename_format', '{published} {creator_name} {id}')
        
        start_page_text = getattr(self.parent_app, 'start_page_input', None) and self.parent_app.start_page_input.text().strip()
        end_page_text = getattr(self.parent_app, 'end_page_input', None) and self.parent_app.end_page_input.text().strip()
        
        self.start_page = 1
        self.end_page = None
        if start_page_text and start_page_text.isdigit():
            self.start_page = max(1, int(start_page_text))
        if end_page_text and end_page_text.isdigit():
            self.end_page = max(1, int(end_page_text))
        
        is_single_post_mode = self.post_id or '/post-' in self.start_url
        album_path = ""
        all_extracted_posts = []
        
        try:
            if is_single_post_mode:
                self.progress_signal.emit("   Mode: Single Post detected.")
                album_title, jobs, _, extracted_posts = fetch_single_simpcity_page(self.start_url, self._log_interceptor, cookies=self.cookies, post_id=self.post_id, check_pause_func=self._check_pause_cancel)
                self._collect_external_links(jobs)
                all_extracted_posts.extend(extracted_posts)
                album_path = os.path.join(self.output_dir, clean_folder_name(album_title or "simpcity_post"))
            else:
                self.progress_signal.emit(f"   Mode: Full Thread detected (Pages {self.start_page} to {self.end_page or 'End'}).")
                base_url_for_first = re.sub(r'(/page-\d+)|(/post-\d+)', '', self.start_url).split('#')[0].strip('/')
                first_page_url = base_url_for_first if self.start_page == 1 else f"{base_url_for_first}/page-{self.start_page}"
                album_title, jobs, _, extracted_posts = fetch_single_simpcity_page(first_page_url, self._log_interceptor, cookies=self.cookies, check_pause_func=self._check_pause_cancel)
                self._collect_external_links(jobs)
                all_extracted_posts.extend(extracted_posts)
                album_path = os.path.join(self.output_dir, clean_folder_name(album_title or "simpcity_album"))
                
            if self._check_pause_cancel():
                self.finished_signal.emit(0, 0, True, [])
                return
                
            os.makedirs(album_path, exist_ok=True)
            self.progress_signal.emit(f"   Saving all content to folder: '{os.path.basename(album_path)}'")
        except Exception as e:
            self.progress_signal.emit(f"❌ Could not process the initial page. Aborting. Error: {e}")
            self.finished_signal.emit(0, 0, self.is_cancelled, []); return
            
        num_service_threads = 4  
        service_executor = ThreadPoolExecutor(max_workers=num_service_threads, thread_name_prefix='SimpCityService')
        for _ in range(num_service_threads): 
            service_executor.submit(self._service_worker, album_path)
            
        num_image_threads = 15
        image_executor = ThreadPoolExecutor(max_workers=num_image_threads, thread_name_prefix='SimpCityImage')
        for _ in range(num_image_threads): 
            image_executor.submit(self._image_worker, album_path)

        try:
            if is_single_post_mode:
                _, jobs, _, _ = fetch_single_simpcity_page(self.start_url, self._log_interceptor, cookies=self.cookies, post_id=self.post_id, check_pause_func=self._check_pause_cancel)
                enriched_jobs = self._get_enriched_jobs(jobs)
                if enriched_jobs and not self.is_cancelled:
                    for job in enriched_jobs:
                        if job['type'] == 'image': 
                            if self.should_dl_images: self.image_queue.put(job)
                        else: self.service_queue.put(job)         
         
            else:
                base_url = re.sub(r'(/page-\d+)|(/post-\d+)', '', self.start_url).split('#')[0].strip('/')
                page_counter = self.start_page; end_of_thread = False; MAX_RETRIES = 3
                while not end_of_thread:
                    if self.end_page and page_counter > self.end_page:
                        self.progress_signal.emit(f"   -> Reached target end page ({self.end_page}). Stopping crawl.")
                        break

                    if self._check_pause_cancel(): break
                    page_url = f"{base_url}/page-{page_counter}" if page_counter > 1 else base_url
                    retries = 0; page_fetch_successful = False
                    while retries < MAX_RETRIES:
                        if self._check_pause_cancel(): end_of_thread = True; break
                        self.progress_signal.emit(f"\n--- Analyzing page {page_counter} (Attempt {retries + 1}/{MAX_RETRIES}) ---")
                        try:
                            page_title, jobs_on_page, final_url, extracted_posts = fetch_single_simpcity_page(page_url, self._log_interceptor, cookies=self.cookies, check_pause_func=self._check_pause_cancel)
                            if page_counter > self.start_page:
                                self._collect_external_links(jobs_on_page)
                                all_extracted_posts.extend(extracted_posts)
                            
                            if self.is_cancelled: end_of_thread = True; break
                            
                            if final_url != page_url:
                                self.progress_signal.emit(f"   -> Redirect detected from {page_url} to {final_url}")
                                try:
                                    req_page_match = re.search(r'/page-(\d+)', page_url)
                                    final_page_match = re.search(r'/page-(\d+)', final_url)

                                    if req_page_match:
                                        req_page_num = int(req_page_match.group(1))

                                        if final_page_match and int(final_page_match.group(1)) < req_page_num:
                                            self.progress_signal.emit(f"   -> Redirected to an earlier page ({final_page_match.group(0)}). Reached end of thread.")
                                            end_of_thread = True
                                        
                                        elif not final_page_match and req_page_num > 1:
                                            self.progress_signal.emit(f"   -> Redirected to base thread URL. Reached end of thread.")
                                            end_of_thread = True

                                except (ValueError, TypeError):
                                    pass
                            
                            if end_of_thread:
                                page_fetch_successful = True; break

                            if page_counter > self.start_page and not page_title:
                                self.progress_signal.emit(f"   -> Page {page_counter} is invalid or has no title. Reached end of thread.")
                                end_of_thread = True
                            elif not jobs_on_page: 
                                self.progress_signal.emit(f"   -> Page {page_counter} has no content. Reached end of thread.")
                                end_of_thread = True
                            else:
                                new_jobs = [job for job in jobs_on_page if job.get('url') not in self.processed_job_urls]
                                if not new_jobs and page_counter > self.start_page: 
                                    self.progress_signal.emit(f"   -> Page {page_counter} contains no new content. Reached end of thread.")
                                    end_of_thread = True
                                else:
                                    enriched_jobs = self._get_enriched_jobs(new_jobs)
                                    if not enriched_jobs and not new_jobs:
                                        self.progress_signal.emit(f"   -> Page {page_counter} content was filtered out. Reached end of thread.")
                                        end_of_thread = True

                                    else:
                                        for job in enriched_jobs:
                                            self.processed_job_urls.add(job.get('url'))
                                            if job['type'] == 'image':
                                                if self.should_dl_images: self.image_queue.put(job)
                                            else: self.service_queue.put(job)

                            page_fetch_successful = True; break
                        except requests.exceptions.HTTPError as e:
                            if e.response.status_code in [403, 404]: 
                                self.progress_signal.emit(f"   -> Page {page_counter} returned {e.response.status_code}. Reached end of thread.")
                                end_of_thread = True; break
                            elif e.response.status_code == 429: 
                                self.progress_signal.emit(f"   -> Rate limited (429). Waiting...")
                                if self._smart_sleep(5 * (retries + 2)):
                                    end_of_thread = True
                                    break
                                retries += 1
                            else: 
                                self.progress_signal.emit(f"   -> HTTP Error {e.response.status_code} on page {page_counter}. Stopping crawl.")
                                end_of_thread = True; break
                        except Exception as e:
                            self.progress_signal.emit(f"   Stopping crawl due to error on page {page_counter}: {e}"); end_of_thread = True; break
                    if not page_fetch_successful and not end_of_thread: 
                        self.progress_signal.emit(f"   -> Failed to fetch page {page_counter} after {MAX_RETRIES} attempts. Stopping crawl.")
                        end_of_thread = True
                    if not end_of_thread: page_counter += 1
        except Exception as e:
            self.progress_signal.emit(f"❌ A critical error occurred during the main fetch phase: {e}")

        self.progress_signal.emit("\n--- All pages analyzed. Waiting for background downloads to complete... ---")
        
        if self.should_extract_text and all_extracted_posts:
            self._export_text(all_extracted_posts, album_title, album_path)
            
        self._save_external_links()

        for _ in range(num_image_threads): self.image_queue.put(None)
        for _ in range(num_service_threads): self.service_queue.put(None)
        
        image_executor.shutdown(wait=True)
        service_executor.shutdown(wait=True)
        
        self.finished_signal.emit(self.total_dl_count, self.total_skip_count, self.is_cancelled, [])

    def _save_external_links(self):
        config = self.external_links_config
        if not config or not self.collected_external_links: return
        
        filepath = config.get('filepath')
        if not filepath:
            filepath = os.path.join(self.output_dir, f"external_links.{config.get('format', 'txt')}")
            
        import csv, json
        
        if config.get('separate_files'):
            links_by_platform = {}
            for item in self.collected_external_links:
                plat = item.get('platform', 'unknown')
                sanitized = re.sub(r'[<>:"/\\|?*]', '_', plat.lower().replace(' ', '_'))
                links_by_platform.setdefault(sanitized, []).append(item)
                
            base, ext = os.path.splitext(filepath)
            if not ext: ext = f".{config.get('format', 'txt')}"
            
            for plat, items in links_by_platform.items():
                plat_filepath = f"{base}_{plat}{ext}"
                self._write_links_format(plat_filepath, items, config)
        else:
            self._write_links_format(filepath, self.collected_external_links, config)
            
        self.progress_signal.emit(f"   ✅ Saved {len(self.collected_external_links)} external links.")

    def _write_links_format(self, filepath, items, config):
        fmt = config.get('format', 'txt')
        mode = 'a' if os.path.exists(filepath) else 'w'
        import json, csv
        
        if fmt == 'json':
            data = []
            if os.path.exists(filepath):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                except: pass
            data.extend(items)
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            return
            
        with open(filepath, mode, encoding='utf-8', newline='') as f:
            if fmt == 'csv':
                writer = csv.DictWriter(f, fieldnames=["url", "platform", "post_title", "thread_title", "link_text", "key", "post_id", "username", "date"])
                if mode == 'w': writer.writeheader()
                for item in items:
                    writer.writerow({k: v for k, v in item.items() if k in writer.fieldnames})
                return
                
            is_md = (fmt == 'md')
            mode_txt = config.get('txt_mode', 'simple')
            
            if mode_txt == 'simple':
                for item in items:
                    if is_md:
                        f.write(f"- [{item.get('link_text') or item.get('url')}]({item.get('url')})\n")
                    else:
                        f.write(item.get('url', '') + "\n")
            elif mode_txt == 'detailed':
                current_title = None
                for item in items:
                    title = item.get('thread_title') or item.get('post_title') or 'Unknown'
                    if config.get('include_titles') and title != current_title:
                        if current_title is not None:
                            f.write("\n" + ("="*60 if not is_md else "---") + "\n\n")
                        if is_md: f.write(f"### {title}\n")
                        else: f.write(f"# Post/Thread: {title}\n")
                        current_title = title
                        
                    if is_md:
                        line = f"- **URL:** [{item.get('link_text') or item.get('url')}]({item.get('url')})"
                        if config.get('include_platform') and item.get('platform'): line += f" | **Platform:** {item['platform']}"
                        f.write(line + "\n")
                    else:
                        line_parts = [item.get('url', '')]
                        if config.get('include_platform') and item.get('platform'): line_parts.append(f"Platform: {item['platform']}")
                        if config.get('include_link_text') and item.get('link_text'): line_parts.append(f"Description: {item['link_text']}")
                        f.write(" | ".join(line_parts) + "\n")
            elif mode_txt == 'custom':
                template = config.get('custom_template', '{url}')
                try:
                    from src.ui.dialogs.ExportLinksDialog import ExportLinksDialog
                    key_extractor = ExportLinksDialog.extract_key_from_url
                except ImportError:
                    key_extractor = lambda u: ""
                    
                for item in items:
                    try:
                        url_val = item.get('url', '')
                        key_val = item.get('key', '')
                        if not key_val:
                            key_val = key_extractor(url_val)
                            
                        formatted_line = template.format(
                            url=url_val,
                            post_title=item.get('post_title', ''),
                            link_text=item.get('link_text', ''),
                            platform=item.get('platform', ''),
                            key=key_val,
                            thread_title=item.get('thread_title', ''),
                            post_id=item.get('post_id', ''),
                            username=item.get('username', ''),
                            date=item.get('date', '')
                        )
                    except KeyError:
                        formatted_line = template
                    f.write(formatted_line)
                    if not template.endswith('\n'):
                        f.write('\n')

    def _export_text(self, posts: list, album_title: str, album_path: str):
        format_choice = self.text_export_format.lower()
        single_pdf_mode = getattr(self.parent_app, 'single_pdf_setting', False)
        
        unique_posts = []
        seen_ids = set()
        for post in posts:
            post_id = str(post.get('post_id'))
            if post_id and post_id != "unknown" and post_id in seen_ids:
                continue
            if post_id and post_id != "unknown":
                seen_ids.add(post_id)
            unique_posts.append(post)

        if format_choice == 'pdf' and single_pdf_mode:
            all_posts_data = []
            font_path = ""
            project_root_dir = getattr(self.parent_app, 'app_base_dir', '')
            if project_root_dir:
                base_dir = getattr(sys, '_MEIPASS', project_root_dir)
                font_path = os.path.join(base_dir, 'data', 'dejavu-sans', 'DejaVuSans.ttf')
                
            for post in unique_posts:
                post_data = {
                    'title': post.get('creator_name', 'SimpCity Post'),
                    'published': post.get('published', 'Unknown'),
                    'creator_name': post.get('creator_name', 'Unknown'),
                    'service': post.get('service', 'Unknown'),
                    'original_link': f"{self.start_url.split('#')[0]}/post-{post.get('post_id', '')}",
                    'content': post.get('content', '')
                }
                all_posts_data.append(post_data)

            safe_title = "".join([c for c in (album_title or "SimpCity_Thread") if c.isalpha() or c.isdigit() or c==' ']).rstrip()
            final_save_path = os.path.join(album_path, f"{safe_title}_FullThread.pdf")
            self.progress_signal.emit(f"   [SimpCity] 📄 Exporting single PDF containing {len(unique_posts)} posts: {os.path.basename(final_save_path)}")
            
            create_single_pdf_from_content(
                posts_data=all_posts_data,
                output_filename=final_save_path,
                font_path=font_path,
                add_info_page=self.add_info_in_pdf_setting,
                continuous=True,
                logger=self.progress_signal.emit
            )
            return

        self.progress_signal.emit(f"\n--- Generating {format_choice.upper()} for extracted text ({len(unique_posts)} posts) ---")
            
        posts_by_id = {str(post.get('post_id')): post for post in unique_posts if post.get('post_id')}
        top_level_posts = []
        
        for post in unique_posts:
            reply_id = post.get('reply_to_post_id')
            if reply_id and str(reply_id) in posts_by_id:
                parent_post = posts_by_id[str(reply_id)]
                reply_text = post.get('content', '').strip()
                if reply_text:
                    append_str = (f"\n\n{'='*40}\n"
                                  f"Reply by {post.get('creator_name', 'Unknown')} on {post.get('published', 'Unknown')}:\n"
                                  f"{'-'*40}\n{reply_text}")
                    parent_post['content'] = parent_post.get('content', '') + append_str
            else:
                top_level_posts.append(post)
        
        for post in top_level_posts:
            if not post.get('content') or not post['content'].strip():
                continue

            base_name = f"post_{post.get('post_id')}.{format_choice}"
            base_name = self._apply_simpcity_custom_renaming(base_name, post)
            
            target_album_path = self._get_target_path(album_path, post)
            final_save_path = os.path.join(target_album_path, base_name)
            
            base, ext = os.path.splitext(final_save_path)
            counter = 1
            while os.path.exists(final_save_path):
                final_save_path = f"{base}_{counter}{ext}"
                counter += 1

            try:
                if format_choice == 'pdf':
                    font_path = ""
                    project_root_dir = getattr(self.parent_app, 'app_base_dir', '')
                    if project_root_dir:
                        base_dir = getattr(sys, '_MEIPASS', project_root_dir)
                        font_path = os.path.join(base_dir, 'data', 'dejavu-sans', 'DejaVuSans.ttf')
                    
                    post_data_for_pdf = {
                        'title': post.get('creator_name', 'SimpCity Post'),
                        'published': post.get('published', 'Unknown'),
                        'creator_name': post.get('creator_name', 'Unknown'),
                        'service': post.get('service', 'Unknown'),
                        'original_link': f"{self.start_url.split('#')[0]}/post-{post.get('post_id', '')}",
                        'content_text_for_pdf': post.get('content', '')
                    }
                    create_individual_pdf(
                        post_data=post_data_for_pdf,
                        output_filename=final_save_path,
                        font_path=font_path,
                        add_info_page=True,
                        logger=self.progress_signal.emit
                    )
                elif format_choice == 'docx':
                    if Document:
                        document = Document()
                        document.add_heading(post.get('creator_name', 'SimpCity Post'), 0)
                        document.add_paragraph(f"Date: {post.get('published', 'Unknown')}")
                        document.add_paragraph(f"Post ID: {post.get('post_id', 'Unknown')}")
                        document.add_paragraph(f"URL: {self.start_url.split('#')[0]}/post-{post.get('post_id', '')}")
                        document.add_page_break()
                        document.add_paragraph(post.get('content', ''))
                        document.save(final_save_path)
                    else:
                        self.progress_signal.emit(f"   ⚠️ 'python-docx' library not installed. Saving as .txt.")
                        final_save_path = os.path.splitext(final_save_path)[0] + ".txt"
                        format_choice = 'txt'
                
                if format_choice == 'txt':
                    content_to_write = post.get('content', '')
                    header = (f"Poster: {post.get('creator_name', 'Unknown')}\n"
                              f"Date: {post.get('published', 'Unknown')}\n"
                              f"Post ID: {post.get('post_id', 'Unknown')}\n"
                              f"URL: {self.start_url.split('#')[0]}/post-{post.get('post_id', '')}\n"
                              f"{'-'*40}\n\n")
                    content_to_write = header + content_to_write
                    with open(final_save_path, 'w', encoding='utf-8') as f:
                        f.write(content_to_write)
                self.progress_signal.emit(f"   ✅ Saved {format_choice.upper()}: '{os.path.basename(final_save_path)}'")
            except Exception as e:
                self.progress_signal.emit(f"   ❌ Error saving text file: {e}")